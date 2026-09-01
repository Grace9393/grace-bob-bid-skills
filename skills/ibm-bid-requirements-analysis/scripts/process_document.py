#!/usr/bin/env python3
"""
Extract structured content from tender documents (PDF/DOCX)
"""

import sys
import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

def extract_pdf_content(filepath):
    """Extract text and tables from PDF using pdfplumber"""
    import pdfplumber
    
    with pdfplumber.open(filepath) as pdf:
        text_content = []
        tables = []
        
        for page in pdf.pages:
            # Extract text
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
            
            # Extract tables
            page_tables = page.extract_tables()
            if page_tables:
                tables.extend(page_tables)
        
        return {
            'text': '\n'.join(text_content),
            'tables': tables
        }

def extract_docx_content(filepath):
    """Extract text, tables, and structure from DOCX"""
    from docx import Document

    doc = Document(filepath)

    paragraphs = []
    tables = []
    headings = []

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = 1
            match = re.search(r'Heading\\s+(\\d+)$', para.style.name)
            if match:
                level = int(match.group(1))
            headings.append({
                'level': level,
                'text': para.text
            })
        paragraphs.append(para.text)

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text for cell in row.cells])
        tables.append(table_data)

    return {
        'text': '\n'.join(paragraphs),
        'tables': tables,
        'headings': headings
    }

def extract_docx_markdown(filepath):
    """Extract markdown from DOCX using pandoc if available"""
    if not shutil.which("pandoc"):
        return None

    with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        subprocess.run(
            ["pandoc", "--track-changes=all", str(filepath), "-o", str(tmp_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        return tmp_path.read_text()
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

def parse_markdown_headings(markdown_text):
    """Return heading list from markdown text"""
    headings = []
    in_code_block = False
    for line in markdown_text.splitlines():
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
        match = re.match(r'^(#{1,6})\\s+(.+)$', line)
        if not match:
            continue
        level = len(match.group(1))
        text = match.group(2).strip().strip('#').strip()
        if text:
            headings.append({"level": level, "text": text})
    return headings

# UK Government tender section patterns
SECTION_PATTERNS = {
    'background': r'(background|context|introduction|about\s+(?:the|this)|executive\s+summary)',
    'requirements': r'(requirements?|specification|scope\s+of\s+work|deliverables|outputs)',
    'evaluation': r'(evaluation|assessment|award\s+criteria|scoring|marking\s+scheme)',
    'constraints': r'(constraints?|limitations?|dependencies|assumptions)',
    'timelines': r'(timeline|schedule|key\s+dates|milestones|timetable)',
    'budget': r'(budget|pricing|financial|commercial|contract\s+value)',
    'client': r'(about\s+us|our\s+organisation|organisation\s+overview)',
    'technical': r'(technical\s+(?:requirements|specification)|system\s+requirements)',
}

def identify_sections(text, headings=None):
    """Map content to tender sections"""
    sections = {}
    lines = text.split('\n')

    if headings:
        heading_indices = []
        for idx, line in enumerate(lines):
            match = re.match(r'^(#{1,6})\\s+(.+)$', line)
            if match:
                heading_indices.append((idx, match.group(2).strip()))

        if heading_indices:
            for i, (start_idx, heading_text) in enumerate(heading_indices):
                end_idx = heading_indices[i + 1][0] if i + 1 < len(heading_indices) else len(lines)
                section_type = None
                for section_key, pattern in SECTION_PATTERNS.items():
                    if re.search(pattern, heading_text, re.IGNORECASE):
                        section_type = section_key
                        break
                if not section_type:
                    continue
                content_lines = [l.strip() for l in lines[start_idx + 1:end_idx] if l.strip()]
                if content_lines:
                    sections.setdefault(section_type, []).extend(content_lines)

            for section_type in sections:
                sections[section_type] = '\n'.join(sections[section_type])
            if sections:
                return sections

    current_section = None

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Heuristic: treat only short lines without trailing period as headers
        is_candidate_header = len(line_stripped) <= 120 and not line_stripped.endswith('.')
        if is_candidate_header:
            for section_type, pattern in SECTION_PATTERNS.items():
                if re.search(pattern, line_stripped, re.IGNORECASE):
                    current_section = section_type
                    sections.setdefault(current_section, [])
                    break

        if current_section and line_stripped:
            sections[current_section].append(line_stripped)

    for section_type in sections:
        sections[section_type] = '\n'.join(sections[section_type])

    return sections

def process_tender_document(filepath):
    """Main processing function"""
    filepath = Path(filepath)
    
    if not filepath.exists():
        raise FileNotFoundError(f"File not found: {filepath}")
    
    ext = filepath.suffix.lower()
    
    if ext == '.pdf':
        content = extract_pdf_content(filepath)
    elif ext == '.docx':
        markdown_text = extract_docx_markdown(filepath)
        content = extract_docx_content(filepath)
        if markdown_text:
            content['markdown'] = markdown_text
            content['headings'] = parse_markdown_headings(markdown_text)
            content['text'] = markdown_text
    else:
        raise ValueError(f"Unsupported format: {ext}. Supported: .pdf, .docx")
    
    # Identify sections
    sections = identify_sections(content['text'], content.get('headings'))
    
    return {
        'filepath': str(filepath),
        'format': ext,
        'sections': sections,
        'tables': content['tables'],
        'full_text': content['text'],
        'headings': content.get('headings', [])
    }

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python3 process_document.py <path_to_tender_document>")
        sys.exit(1)
    
    filepath = sys.argv[1]
    
    try:
        result = process_tender_document(filepath)
        print(json.dumps(result, indent=2))
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
