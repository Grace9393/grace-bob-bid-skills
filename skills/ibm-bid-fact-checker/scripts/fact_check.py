#!/usr/bin/env python3
"""
IBM Bid Fact Checker Script

This script performs automated fact-checking of tender responses against source documents.
It extracts claims and metrics from the response, cross-references them with source documents,
and generates a comprehensive verification report.

Usage:
    python3 fact_check.py --response <response_file> --sources <source_files>
    python3 fact_check.py --response response.docx --sources rfp.pdf ibm_capabilities.pdf
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List, Dict

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None


class FactChecker:
    """Main fact-checking class"""

    def __init__(self, response_file: str, source_files: List[str]):
        self.response_file = response_file
        self.source_files = source_files
        self.response_text = ""
        self.source_texts = {}
        self.claims = []
        self.verified_claims = []
        self.flagged_claims = []
        self.evidence_gaps = []
        self._token_re = re.compile(r"[a-zA-Z0-9]+")

    def _tokenize(self, text: str) -> List[str]:
        return [t.lower() for t in self._token_re.findall(text or "") if len(t) > 2]

    def _partial_match(self, claim_text: str, source_text: str) -> bool:
        claim_tokens = set(self._tokenize(claim_text))
        if not claim_tokens:
            return False
        source_tokens = set(self._tokenize(source_text))
        overlap = claim_tokens.intersection(source_tokens)
        return len(overlap) >= 3 and (len(overlap) / max(len(claim_tokens), 1)) >= 0.4

    def load_response(self) -> bool:
        """Load the response document"""
        try:
            if self.response_file.endswith(".docx"):
                if not Document:
                    print("Error: python-docx is required for DOCX files")
                    return False
                doc = Document(self.response_file)
                self.response_text = "\n".join([para.text for para in doc.paragraphs])
            elif self.response_file.endswith(".pdf"):
                if not PdfReader:
                    print("Error: PyPDF2 is required for PDF files")
                    return False
                reader = PdfReader(self.response_file)
                text = ""
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
                self.response_text = text
            else:
                with open(self.response_file, "r", encoding="utf-8") as f:
                    self.response_text = f.read()
            print(f"✓ Loaded response document: {self.response_file}")
            return True
        except Exception as e:
            print(f"Error loading response document: {e}")
            return False

    def load_sources(self) -> bool:
        """Load source documents"""
        success = True
        for source_file in self.source_files:
            try:
                if source_file.endswith(".pdf"):
                    if not PdfReader:
                        print(f"Error: PyPDF2 is required for PDF files")
                        success = False
                        continue
                    reader = PdfReader(source_file)
                    text = ""
                    for page in reader.pages:
                        text += (page.extract_text() or "") + "\n"
                    self.source_texts[source_file] = text
                    print(f"✓ Loaded source document: {source_file}")
                elif source_file.endswith(".docx"):
                    if not Document:
                        print(f"Error: python-docx is required for DOCX files")
                        success = False
                        continue
                    doc = Document(source_file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    self.source_texts[source_file] = text
                    print(f"✓ Loaded source document: {source_file}")
                else:
                    with open(source_file, "r", encoding="utf-8") as f:
                        text = f.read()
                    self.source_texts[source_file] = text
                    print(f"✓ Loaded source document: {source_file}")
            except Exception as e:
                print(f"Error loading source document {source_file}: {e}")
                success = False
        return success

    def extract_claims(self) -> List[Dict]:
        """Extract claims from the response document"""
        claims = []

        # Pattern 1: Performance metrics (percentages, numbers)
        metric_patterns = [
            r"(\d+\.?\d*)\s*%?\s*(availability|uptime|performance|throughput|response time)",
            r"(\d+\.?\d*)\s*(SLA|service level)",
            r"(\d+\.?\d*)\s*(availability|uptime)",
        ]

        # Pattern 2: Capability claims
        capability_patterns = [
            r"(IBM can|IBM provides|IBM offers|IBM delivers)\s+(.*?)(?:\.|,|\n)",
            r"(IBM is|IBM has)\s+(.*?)(?:\.|,|\n)",
            r"(We can|We provide|We offer|We deliver)\s+(.*?)(?:\.|,|\n)",
            r"(We have)\s+(.*?)(?:\.|,|\n)",
        ]

        # Pattern 3: Pricing claims
        pricing_patterns = [
            r"\$?(\d+\.?\d*)\s*(million|thousand|billion)",
            r"(cost|price|investment|budget)\s*(?:of|for|to)\s*\$?(\d+\.?\d*)",
        ]

        # Pattern 4: Requirement compliance claims
        compliance_patterns = [
            r"(meets|fulfills|satisfies|addresses)\s+(.*?)(?:\.|,|\n)",
            r"(complies with|adheres to|follows)\s+(.*?)(?:\.|,|\n)",
        ]

        # Combine all patterns
        all_patterns = [
            ("metric", metric_patterns),
            ("capability", capability_patterns),
            ("pricing", pricing_patterns),
            ("compliance", compliance_patterns),
        ]

        for claim_type, patterns in all_patterns:
            for pattern in patterns:
                matches = re.finditer(pattern, self.response_text, re.IGNORECASE)
                for match in matches:
                    claim_text = match.group(0).strip()
                    # Avoid duplicates
                    if not any(c["text"] == claim_text for c in claims):
                        claims.append(
                            {
                                "text": claim_text,
                                "type": claim_type,
                                "confidence": "medium",
                            }
                        )

        # Also extract sentences with strong assertions
        sentence_pattern = r"([^.!?]+[.!?])"
        sentences = re.findall(sentence_pattern, self.response_text)
        for sentence in sentences:
            sentence = sentence.strip()
            # Check for strong assertion words
            assertion_words = [
                "guarantee",
                "guarantees",
                "ensure",
                "ensures",
                "provide",
                "provides",
                "deliver",
                "delivers",
                "achieve",
                "achieves",
                "establish",
                "establishes",
                "maintain",
                "maintains",
                "enable",
                "enables",
                "secure",
                "secures",
                "assure",
                "assures",
                "confirm",
                "confirms",
                "verify",
                "verifies",
                "validate",
                "validates",
                "certify",
                "certifies",
                "prove",
                "proves",
                "demonstrate",
                "demonstrates",
                "realise",
                "realises",
                "fulfil",
                "fulfils",
                "meet",
                "meets",
                "satisfy",
                "satisfies",
                "produce",
                "produces",
                "yield",
                "yields",
                "generate",
                "generates",
                "create",
                "creates",
            ]
            if any(word in sentence.lower() for word in assertion_words):
                if sentence not in [c["text"] for c in claims]:
                    claims.append(
                        {"text": sentence, "type": "assertion", "confidence": "low"}
                    )

        self.claims = claims
        print(f"✓ Extracted {len(claims)} claims from response")
        return claims

    def verify_claim(self, claim: Dict) -> Dict:
        """Verify a single claim against source documents"""
        claim_text = claim["text"].lower()
        sources_found = []

        for source_file, source_text in self.source_texts.items():
            # Search for the claim text in source
            if claim_text in source_text.lower():
                sources_found.append({"file": source_file, "match_type": "exact"})
            else:
                # Try partial matching with token overlap
                if self._partial_match(claim_text, source_text):
                    sources_found.append({"file": source_file, "match_type": "partial"})

        # Determine verification status
        if sources_found:
            return {
                **claim,
                "verified": True,
                "sources": sources_found,
                "evidence_strength": "strong"
                if sources_found[0]["match_type"] == "exact"
                else "medium",
            }
        else:
            return {
                **claim,
                "verified": False,
                "sources": [],
                "evidence_strength": "none",
                "recommendation": "Consider removing or finding supporting evidence",
            }

    def verify_all_claims(self):
        """Verify all claims against source documents"""
        print("\nVerifying claims against source documents...")
        for i, claim in enumerate(self.claims, 1):
            verified_claim = self.verify_claim(claim)
            if verified_claim["verified"]:
                self.verified_claims.append(verified_claim)
            else:
                self.flagged_claims.append(verified_claim)
            print(
                f"  [{i}/{len(self.claims)}] {verified_claim['text'][:60]}... - {'✓' if verified_claim['verified'] else '✗'}"
            )

        print(
            f"\n✓ Verification complete: {len(self.verified_claims)} verified, {len(self.flagged_claims)} flagged"
        )

    def generate_report(self) -> str:
        """Generate a comprehensive fact-check report"""
        report = []
        report.append("# Fact-Check Report")
        report.append("")
        report.append(f"**Response Document:** {self.response_file}")
        report.append(f"**Source Documents:** {', '.join(self.source_files)}")
        report.append(
            f"**Date:** {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        report.append("")
        report.append("---")
        report.append("")

        # Summary
        report.append("## Summary")
        report.append("")
        report.append(f"- Total claims checked: {len(self.claims)}")
        total_claims = len(self.claims)
        verified_count = len(self.verified_claims)
        flagged_count = len(self.flagged_claims)
        verified_pct = (verified_count / total_claims * 100.0) if total_claims else 0.0
        flagged_pct = (flagged_count / total_claims * 100.0) if total_claims else 0.0
        report.append(
            f"- Verified claims: {verified_count} ({verified_pct:.1f}%)"
        )
        report.append(
            f"- Flagged claims: {flagged_count} ({flagged_pct:.1f}%)"
        )
        report.append("")

        # Verified claims
        if self.verified_claims:
            report.append("## Verified Claims")
            report.append("")
            for claim in self.verified_claims:
                report.append(f"### {claim['text'][:100]}...")
                report.append("")
                report.append(
                    f"**Evidence Strength:** {claim['evidence_strength'].upper()}"
                )
                report.append("")
                report.append("**Source Documents:**")
                for source in claim["sources"]:
                    report.append(f"- {source['file']} ({source['match_type']})")
                report.append("")

        # Flagged claims
        if self.flagged_claims:
            report.append("## Flagged Claims")
            report.append("")
            for i, claim in enumerate(self.flagged_claims, 1):
                report.append(f"### {i}. {claim['text'][:100]}...")
                report.append("")
                report.append(
                    f"**Evidence Strength:** {claim['evidence_strength'].upper()}"
                )
                report.append("")
                report.append("**Recommendation:**")
                report.append(f"- {claim['recommendation']}")
                report.append("")

        # Evidence gaps
        if self.evidence_gaps:
            report.append("## Evidence Gaps")
            report.append("")
            for gap in self.evidence_gaps:
                report.append(f"- {gap}")
            report.append("")

        # Recommendations
        report.append("## Recommendations")
        report.append("")
        if self.flagged_claims:
            report.append("1. Review and revise flagged claims")
            report.append("2. Remove unsupported claims or find supporting evidence")
            report.append("3. Add source document citations to verified claims")
            report.append("4. Consider adding a references section to the response")
        else:
            report.append("1. All claims are verified against source documents")
            report.append("2. Consider adding source citations for transparency")
        report.append("")

        return "\n".join(report)

    def save_report(self, report: str, output_file: str = "fact_check_report.md"):
        """Save the fact-check report to a file"""
        try:
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(report)
            print(f"\n✓ Report saved to: {output_file}")
            return True
        except Exception as e:
            print(f"Error saving report: {e}")
            return False

    def run(self, output_file: str = "fact_check_report.md"):
        """Run the complete fact-checking process"""
        print("=" * 60)
        print("IBM Bid Fact Checker")
        print("=" * 60)
        print()

        # Load documents
        if not self.load_response():
            return False
        if not self.load_sources():
            return False

        # Extract claims
        self.extract_claims()

        # Verify claims
        self.verify_all_claims()

        # Generate report
        report = self.generate_report()
        self.save_report(report, output_file)

        print("\n" + "=" * 60)
        print("Fact-checking complete!")
        print("=" * 60)

        return True


def main():
    parser = argparse.ArgumentParser(
        description="IBM Bid Fact Checker - Verify tender responses against source documents"
    )
    parser.add_argument(
        "--response",
        required=True,
        help="Path to the response document (PDF, DOCX, or TXT)",
    )
    parser.add_argument(
        "--sources",
        nargs="+",
        required=True,
        help="One or more source documents (PDF, DOCX, or TXT) to verify against",
    )
    parser.add_argument(
        "--output",
        default="fact_check_report.md",
        help="Output report file (default: fact_check_report.md)",
    )

    args = parser.parse_args()

    # Validate input files
    response_path = Path(args.response)
    if not response_path.exists():
        print(f"Error: Response file not found: {args.response}")
        sys.exit(1)

    source_paths = [Path(s) for s in args.sources]
    for source_path in source_paths:
        if not source_path.exists():
            print(f"Error: Source file not found: {source_path}")
            sys.exit(1)

    # Run fact-checker
    fact_checker = FactChecker(args.response, args.sources)
    success = fact_checker.run(args.output)

    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
